from __future__ import annotations

import math
import re
import tempfile
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path
from typing import BinaryIO, Iterable

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo
from PIL import Image, ImageDraw, ImageFont


ACCENT = "1D1D1F"
BLUE = "0071E3"
LIGHT_BLUE = "EAF4FF"
GREEN = "1F8A5B"
YELLOW = "F2A900"
RED = "D92D20"
GRAY = "F5F5F7"
MID_GRAY = "86868B"


@dataclass
class AnalysisResult:
    final: pd.DataFrame
    sabc_summary: pd.DataFrame
    inventory_status_summary: pd.DataFrame
    action_summary: pd.DataFrame
    matrix: pd.DataFrame
    sales_year_summary: pd.DataFrame | None
    sales_purchase_year_summary: pd.DataFrame | None
    sales_location_summary: pd.DataFrame | None
    sales_year_location_summary: pd.DataFrame | None
    stock_location_summary: pd.DataFrame | None
    purchase_summary: pd.DataFrame | None
    purchase_sku_summary: pd.DataFrame | None
    purchase_location_summary: pd.DataFrame | None
    purchase_year_location_summary: pd.DataFrame | None
    purchase_sku_location_summary: pd.DataFrame | None
    location_year_business_view: pd.DataFrame | None
    insights: dict
    detected_columns: dict


def normalize_header(value: object) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value).strip().lower())


def normalize_sku(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)):
        if math.isnan(float(value)):
            return ""
        return str(int(round(float(value)))) if float(value).is_integer() else f"{float(value):.0f}"

    text = str(value).strip()
    if not text or text.lower() in {"nan", "none"}:
        return ""
    text = text.replace("\u3000", " ").strip()
    text = re.sub(r"\.0$", "", text)
    if re.fullmatch(r"[0-9]+", text):
        return text
    if re.fullmatch(r"[0-9]+(\.[0-9]+)?[eE][+-]?[0-9]+", text) or re.fullmatch(r"[0-9]+\.[0-9]+", text):
        try:
            return str(int(Decimal(text).to_integral_value()))
        except (InvalidOperation, ValueError):
            return text
    return text


def first_existing_column(columns: Iterable[str], aliases: Iterable[str]) -> str | None:
    normalized = {normalize_header(c): c for c in columns}
    for alias in aliases:
        hit = normalized.get(normalize_header(alias))
        if hit is not None:
            return hit
    return None


def detect_column(df: pd.DataFrame, aliases: list[str], required: bool = True, label: str = "") -> str | None:
    hit = first_existing_column(df.columns, aliases)
    if hit is None and required:
        pretty = label or aliases[0]
        raise ValueError(f"Could not detect required column: {pretty}. Available columns: {list(df.columns)}")
    return hit


def read_excel_any(file: str | Path | BinaryIO | BytesIO) -> pd.DataFrame:
    xl = pd.ExcelFile(file)
    frames = []
    for sheet in xl.sheet_names:
        df = pd.read_excel(xl, sheet_name=sheet)
        if not df.empty:
            frames.append(df)
    if not frames:
        raise ValueError("The uploaded Excel file does not contain readable rows.")
    return pd.concat(frames, ignore_index=True)


def read_excel_with_detected_header(file: str | Path | BinaryIO | BytesIO) -> pd.DataFrame:
    xl = pd.ExcelFile(file)
    frames = []
    for sheet in xl.sheet_names:
        raw = pd.read_excel(xl, sheet_name=sheet, header=None)
        if raw.empty:
            continue
        header_idx = detect_header_row(raw)
        if header_idx is None:
            df = pd.read_excel(xl, sheet_name=sheet)
        else:
            df = build_dataframe_from_header(raw, header_idx)
        if not df.empty:
            frames.append(df)
    if not frames:
        raise ValueError("The uploaded Excel file does not contain readable rows.")
    return pd.concat(frames, ignore_index=True)


def read_year_labeled_files(
    file_or_files: str | Path | BinaryIO | BytesIO | list[str | Path | BinaryIO | BytesIO],
    years: list[int] | None = None,
    detected_header: bool = False,
) -> pd.DataFrame:
    files = file_or_files if isinstance(file_or_files, list) else [file_or_files]
    frames = []
    for idx, file in enumerate(files):
        frame = read_excel_with_detected_header(file) if detected_header else read_excel_any(file)
        if years and idx < len(years):
            frame["_source_year"] = years[idx]
        frames.append(frame)
    if not frames:
        raise ValueError("No readable Excel files were uploaded.")
    return pd.concat(frames, ignore_index=True)


def detect_header_row(raw: pd.DataFrame) -> int | None:
    header_tokens = {
        "sku",
        "productsku",
        "productcode",
        "code",
        "jan",
        "barcode",
        "recommendation",
        "category",
        "productname",
        "wholesaleprice",
        "outercase",
    }
    best_idx = None
    best_score = 0
    for idx, row in raw.head(30).iterrows():
        normalized_values = {normalize_header(v) for v in row.dropna().tolist()}
        score = sum(1 for token in header_tokens if token in normalized_values)
        if score > best_score:
            best_score = score
            best_idx = idx
    return best_idx if best_score >= 2 else None


def build_dataframe_from_header(raw: pd.DataFrame, header_idx: int) -> pd.DataFrame:
    header = raw.iloc[header_idx].tolist()
    subheader = raw.iloc[header_idx + 1].tolist() if header_idx + 1 < len(raw) else [None] * len(header)
    header_non_empty = sum(pd.notna(v) for v in header)
    subheader_non_empty = sum(pd.notna(v) for v in subheader)
    use_subheader = 0 < subheader_non_empty <= max(3, header_non_empty * 0.5)
    columns = []
    last_parent = ""
    for parent, child in zip(header, subheader):
        parent_text = "" if pd.isna(parent) else str(parent).replace("\n", " ").strip()
        child_text = "" if (not use_subheader or pd.isna(child)) else str(child).replace("\n", " ").strip()
        if parent_text:
            last_parent = parent_text
        if parent_text and child_text:
            name = f"{parent_text} {child_text}"
        elif parent_text:
            name = parent_text
        elif child_text and last_parent:
            name = f"{last_parent} {child_text}"
        elif child_text:
            name = child_text
        else:
            name = f"Column {len(columns) + 1}"
        columns.append(name)
    df = raw.iloc[header_idx + (2 if use_subheader else 1) :].copy()
    df.columns = dedupe_columns(columns)
    df = df.dropna(how="all")
    return df


def dedupe_columns(columns: list[str]) -> list[str]:
    seen = {}
    output = []
    for col in columns:
        base = col.strip() or "Column"
        count = seen.get(base, 0)
        seen[base] = count + 1
        output.append(base if count == 0 else f"{base}.{count + 1}")
    return output


def latest_12_months(df: pd.DataFrame, date_col: str | None) -> pd.DataFrame:
    if date_col is None:
        return df.copy()
    work = df.copy()
    work[date_col] = pd.to_datetime(work[date_col], errors="coerce")
    valid_dates = work[date_col].dropna()
    if valid_dates.empty:
        return work
    end_month = valid_dates.max().to_period("M").to_timestamp()
    start_month = (end_month.to_period("M") - 11).to_timestamp()
    month_values = work[date_col].dt.to_period("M").dt.to_timestamp()
    return work[(month_values >= start_month) & (month_values <= end_month)].copy()


def mode_or_first(series: pd.Series) -> str:
    cleaned = series.dropna().astype(str).str.strip()
    cleaned = cleaned[cleaned != ""]
    if cleaned.empty:
        return ""
    mode = cleaned.mode()
    return mode.iloc[0] if not mode.empty else cleaned.iloc[0]


def filter_by_keyword(df: pd.DataFrame, keyword: str, columns: list[str | None]) -> pd.DataFrame:
    cleaned = keyword.strip()
    if not cleaned:
        return df.copy()
    text_cols = [col for col in columns if col and col in df.columns]
    if not text_cols:
        return df.copy()
    mask = pd.Series(False, index=df.index)
    for col in text_cols:
        mask = mask | df[col].astype(str).str.contains(cleaned, case=False, na=False)
    return df[mask].copy()


def clean_location(value: object) -> str:
    if pd.isna(value):
        return "Unassigned"
    text = str(value).strip()
    return text if text else "Unassigned"


def filter_by_location(df: pd.DataFrame, location_col: str | None, location_filter: str = "") -> pd.DataFrame:
    cleaned = location_filter.strip()
    if not cleaned or not location_col or location_col not in df.columns:
        return df.copy()
    mask = df[location_col].astype(str).str.contains(cleaned, case=False, na=False)
    return df[mask].copy()


def aggregate_sales(
    sales_df: pd.DataFrame,
    filter_keyword: str = "",
    manual_year_mode: bool = False,
    location_filter: str = "",
) -> tuple[pd.DataFrame, pd.DataFrame | None, pd.DataFrame | None, pd.DataFrame | None, dict]:
    sku_col = detect_column(sales_df, ["Product SKU", "SKU", "Product Code", "Item SKU", "Barcode"], label="Product SKU")
    name_col = detect_column(sales_df, ["Product / Service", "Product Name", "Name", "Item Name", "Description"], label="Product Name")
    qty_col = detect_column(sales_df, ["Quantity", "Qty", "Units", "Units Sold", "Sold Qty"], label="Quantity")
    date_col = detect_column(sales_df, ["Month, Year", "Month/Year", "Month Year", "Date", "Order Date"], required=False)
    location_col = detect_column(sales_df, ["Location", "Warehouse", "Store", "Branch", "Sales Location", "Stock Location", "Outlet"], required=False)
    amount_col = detect_column(sales_df, ["Total without taxes", "Net Sales", "Sales Amount", "Amount", "Revenue", "Total with taxes", "Total"], required=False)
    barcode_col = detect_column(sales_df, ["Barcode", "JAN", "UPC", "EAN"], required=False)
    brand_col = detect_column(sales_df, ["Brand", "Supplier", "Vendor"], required=False)

    filtered = filter_by_keyword(sales_df, filter_keyword, [name_col, sku_col, barcode_col, brand_col])
    filtered = filter_by_location(filtered, location_col, location_filter)
    work = filtered.copy() if manual_year_mode and "_source_year" in filtered.columns else latest_12_months(filtered, date_col)
    work["_sku"] = work[sku_col].map(normalize_sku)
    work["_qty"] = pd.to_numeric(work[qty_col], errors="coerce").fillna(0)
    work = work[work["_sku"] != ""]
    manual_location = location_filter.strip() if location_filter.strip() and not location_col else ""
    grouped = (
        work.groupby("_sku", as_index=False)
        .agg(**{"Product SKU": ("_sku", "first"), "Product Name": (name_col, mode_or_first), "Qty": ("_qty", "sum")})
        .drop(columns=["_sku"], errors="ignore")
        .sort_values("Qty", ascending=False)
    )
    sales_year_summary = None
    annual = None
    if manual_year_mode and "_source_year" in filtered.columns:
        annual = filtered.copy()
        annual["_year"] = pd.to_numeric(annual["_source_year"], errors="coerce")
    elif date_col:
        annual = filtered.copy()
        annual[date_col] = pd.to_datetime(annual[date_col], errors="coerce")
        annual["_year"] = annual[date_col].dt.year

    if annual is not None:
        annual["_sku"] = annual[sku_col].map(normalize_sku)
        annual["_qty"] = pd.to_numeric(annual[qty_col], errors="coerce").fillna(0)
        annual = annual[(annual["_sku"] != "") & annual["_year"].notna()]
        aggregations = {"SKU Count": ("_sku", "nunique"), "Sales Qty": ("_qty", "sum")}
        if amount_col:
            annual["_sales_amount"] = pd.to_numeric(annual[amount_col], errors="coerce").fillna(0)
            aggregations["Sales Amount"] = ("_sales_amount", "sum")
        sales_year_summary = (
            annual.groupby("_year", dropna=False)
            .agg(**aggregations)
            .reset_index()
            .rename(columns={"_year": "Year"})
            .sort_values("Year")
        )
        sales_year_summary["Year"] = sales_year_summary["Year"].astype(int)

    sales_location_summary = None
    if location_col or manual_location:
        loc = work.copy()
        loc["_location"] = manual_location or loc[location_col].map(clean_location)
        aggregations = {"SKU Count": ("_sku", "nunique"), "Qty": ("_qty", "sum")}
        if amount_col:
            loc["_sales_amount"] = pd.to_numeric(loc[amount_col], errors="coerce").fillna(0)
            aggregations["Sales Amount"] = ("_sales_amount", "sum")
        sales_location_summary = (
            loc.groupby("_location", dropna=False)
            .agg(**aggregations)
            .reset_index()
            .rename(columns={"_location": "Location"})
            .sort_values("Qty", ascending=False)
        )

    sales_year_location_summary = None
    if annual is not None and (location_col or manual_location):
        annual = annual.copy()
        annual["_location"] = manual_location or annual[location_col].map(clean_location)
        aggregations = {"SKU Count": ("_sku", "nunique"), "Sales Qty": ("_qty", "sum")}
        if amount_col:
            if "_sales_amount" not in annual.columns:
                annual["_sales_amount"] = pd.to_numeric(annual[amount_col], errors="coerce").fillna(0)
            aggregations["Sales Amount"] = ("_sales_amount", "sum")
        sales_year_location_summary = (
            annual.groupby(["_year", "_location"], dropna=False)
            .agg(**aggregations)
            .reset_index()
            .rename(columns={"_year": "Year", "_location": "Location"})
            .sort_values(["Year", "Sales Qty"], ascending=[False, False])
        )
        sales_year_location_summary["Year"] = sales_year_location_summary["Year"].astype(int)

    return grouped, sales_year_summary, sales_location_summary, sales_year_location_summary, {
        "sales_sku": sku_col,
        "sales_name": name_col,
        "sales_qty": qty_col,
        "sales_date": date_col or "Not detected",
        "sales_location": location_col or "Not detected",
        "sales_manual_location_label": manual_location or "Not used",
        "sales_amount": amount_col or "Not detected",
        "sales_filter_keyword": filter_keyword or "Not used",
        "sales_year_mode": "Manual upload year" if manual_year_mode else "Latest 12 months / detected from sales date",
    }


def aggregate_stock(stock_df: pd.DataFrame, filter_keyword: str = "", location_filter: str = "") -> tuple[pd.DataFrame, pd.DataFrame | None, dict]:
    sku_col = detect_column(stock_df, ["Product SKU", "SKU", "Product Code", "Item SKU", "Barcode"], label="SKU")
    name_col = detect_column(stock_df, ["Product / Service", "Product Name", "Name", "Item Name", "Description"], required=False)
    available_col = detect_column(stock_df, ["Available", "Available Qty", "Available Stock"], required=False)
    incoming_col = detect_column(stock_df, ["Incoming", "Incoming Qty", "On Order"], required=False)
    on_hand_col = detect_column(stock_df, ["On Hand", "On hand", "Onhand", "Stock On Hand"], required=False)
    location_col = detect_column(stock_df, ["Location", "Warehouse", "Store", "Branch", "Stock Location", "Receiving Location", "Outlet"], required=False)
    barcode_col = detect_column(stock_df, ["Barcode", "JAN", "UPC", "EAN"], required=False)
    brand_col = detect_column(stock_df, ["Brand", "Supplier", "Vendor"], required=False)

    work = filter_by_keyword(stock_df, filter_keyword, [name_col, sku_col, barcode_col, brand_col])
    work = filter_by_location(work, location_col, location_filter)
    work["_sku"] = work[sku_col].map(normalize_sku)
    work = work[work["_sku"] != ""]
    for col in [available_col, incoming_col, on_hand_col]:
        if col:
            work[col] = pd.to_numeric(work[col], errors="coerce").fillna(0)

    aggregations = {"Product SKU": ("_sku", "first")}
    if name_col:
        aggregations["Stock Product Name"] = (name_col, mode_or_first)
    aggregations["Available"] = (available_col, "sum") if available_col else ("_sku", lambda _: 0)
    aggregations["Incoming"] = (incoming_col, "sum") if incoming_col else ("_sku", lambda _: 0)
    aggregations["On Hand"] = (on_hand_col, "sum") if on_hand_col else ("_sku", lambda _: 0)
    grouped = work.groupby("_sku", as_index=False).agg(**aggregations).drop(columns=["_sku"], errors="ignore")

    stock_location_summary = None
    if location_col:
        loc = work.copy()
        loc["_location"] = loc[location_col].map(clean_location)
        stock_location_summary = (
            loc.groupby("_location", dropna=False)
            .agg(
                **{
                    "SKU Count": ("_sku", "nunique"),
                    "Available": (available_col, "sum") if available_col else ("_sku", lambda _: 0),
                    "Incoming": (incoming_col, "sum") if incoming_col else ("_sku", lambda _: 0),
                    "On Hand": (on_hand_col, "sum") if on_hand_col else ("_sku", lambda _: 0),
                }
            )
            .reset_index()
            .rename(columns={"_location": "Location"})
        )
        stock_location_summary["Future Inventory"] = stock_location_summary["Available"] + stock_location_summary["Incoming"]
        stock_location_summary = stock_location_summary.sort_values("Future Inventory", ascending=False)

    return grouped, stock_location_summary, {
        "stock_sku": sku_col,
        "stock_name": name_col or "Not detected",
        "stock_available": available_col or "Not detected",
        "stock_incoming": incoming_col or "Not detected",
        "stock_on_hand": on_hand_col or "Not detected",
        "stock_location": location_col or "Not detected",
        "stock_location_filter": location_filter or "Not used",
        "stock_filter_keyword": filter_keyword or "Not used",
    }


def apply_catalogue(final: pd.DataFrame, catalogue_df: pd.DataFrame | None) -> tuple[pd.DataFrame, dict]:
    if catalogue_df is None:
        return final, {}
    sku_col = detect_column(catalogue_df, ["Product SKU", "SKU", "JAN", "Barcode", "Product Code", "Product Code Product Code", "Code"], label="Catalogue SKU")
    cat = catalogue_df.copy()
    cat["_sku"] = cat[sku_col].map(normalize_sku)
    optional_map = {
        "Category": ["Category", "Product Category", "Collection"],
        "Recommendation / Lifecycle Status": ["Recommendation / Lifecycle Status", "Lifecycle Status", "Recommendation", "Status"],
        "RP USD": ["RP USD", "Retail Price", "Retail USD", "MSRP"],
        "Wholesale Price": ["Wholesale Price", "Wholesale price (USD)", "Wholesale", "WSP", "Cost"],
        "Wholesale Price JPY": ["Wholesale price (JPY)", "Wholesale Price JPY", "Wholesale price JPY"],
        "Inner Case": ["Inner Case", "Quantity Inner case", "Inner", "Inner Qty"],
        "Outer Case": ["Outer Case", "Quantity Outer case", "Outer", "Outer Qty"],
    }
    detected = {"catalogue_sku": sku_col}
    keep = ["_sku"]
    for out_col, aliases in optional_map.items():
        source = detect_column(cat, aliases, required=False)
        detected[f"catalogue_{out_col}"] = source or "Not detected"
        if source:
            cat[out_col] = cat[source]
            keep.append(out_col)
    cat = cat[keep].drop_duplicates("_sku")
    merged = final.merge(cat, left_on="Product SKU", right_on="_sku", how="left").drop(columns=["_sku"], errors="ignore")
    if {"RP USD", "Wholesale Price"}.issubset(merged.columns):
        rp = pd.to_numeric(merged["RP USD"], errors="coerce")
        wholesale = pd.to_numeric(merged["Wholesale Price"], errors="coerce")
        merged["Margin %"] = np.where(rp > 0, (rp - wholesale) / rp, np.nan)
    if "Outer Case" in merged.columns:
        outer = pd.to_numeric(merged["Outer Case"], errors="coerce")
        merged["MOQ risk"] = np.where((merged["Avg Monthly Sales"] > 0) & (outer > merged["Avg Monthly Sales"] * 3), "High MOQ Risk", "")
    return merged, detected


def classify_lifecycle_status(value: object) -> str:
    text = "" if pd.isna(value) else str(value).strip()
    lower = text.lower()
    compact = re.sub(r"\s+", "", lower)
    if not compact:
        return ""
    if any(token in compact for token in ["停产", "discontinued", "discontinue", "phaseout", "phase-out", "廃番"]):
        return "Discontinued / Phase Out"
    if any(token in compact for token in ["renewal", "renew", "更新", "改版", "新版"]):
        return "Renewal"
    if ("展架" in compact or "display" in compact or "tester" in compact) and any(token in compact for token in ["少", "low", "short", "不足"]):
        return "Display Low Stock"
    if ("展架" in compact or "display" in compact or "tester" in compact) and any(token in compact for token in ["无", "none", "no", "outofstock", "out-of-stock", "缺"]):
        return "Display No Stock"
    if any(token in compact for token in ["new", "comingsoon", "coming", "新品"]):
        return "New / Coming Soon"
    return "Active / Review"


def lifecycle_action(row: pd.Series) -> str:
    lifecycle = row.get("Lifecycle Type", "")
    qty = float(row.get("Qty", 0) or 0)
    sabc = str(row.get("SABC Type", ""))
    inventory_status = str(row.get("Inventory Status", ""))
    if lifecycle == "New / Coming Soon":
        return "Launch Plan / Initial PO"
    if lifecycle == "Discontinued / Phase Out":
        return "Stop PO / Phase Out" if qty > 0 else "Do Not Launch / Archive"
    if lifecycle == "Display Low Stock":
        return "Request Display Support" if sabc in {"S", "A"} else "Review Display Need"
    if lifecycle == "Display No Stock":
        return "Urgent Display Support" if sabc in {"S", "A"} else "Confirm Display Strategy"
    if lifecycle == "Renewal":
        return "Version Transition Plan"
    return str(row.get("Action", "")) or inventory_status


def lifecycle_note(row: pd.Series) -> str:
    lifecycle = row.get("Lifecycle Type", "")
    qty = float(row.get("Qty", 0) or 0)
    if lifecycle == "New / Coming Soon":
        return "New SKU: no historical sales is expected; evaluate launch quantity, channel fit, and display support."
    if lifecycle == "Discontinued / Phase Out" and qty > 0:
        return "Historical sales exists but catalogue indicates discontinued/phase-out; avoid replenishment and plan replacement or sell-through."
    if lifecycle == "Discontinued / Phase Out":
        return "Catalogue indicates discontinued/phase-out; keep out of future launch or replenishment planning."
    if lifecycle == "Display Low Stock":
        return "Display/tester stock is limited; sales may be constrained by weak shelf presence."
    if lifecycle == "Display No Stock":
        return "Display/tester stock is missing; request manufacturer support before channel push."
    if lifecycle == "Renewal":
        return "Renewal item: manage old/new version transition and avoid duplicated inventory."
    return ""


def apply_lifecycle_strategy(final: pd.DataFrame) -> pd.DataFrame:
    if "Recommendation / Lifecycle Status" not in final.columns:
        return final
    work = final.copy()
    work["Lifecycle Type"] = work["Recommendation / Lifecycle Status"].map(classify_lifecycle_status)
    work["Lifecycle Note"] = work.apply(lifecycle_note, axis=1)
    work["Lifecycle Action"] = work.apply(lifecycle_action, axis=1)
    override_mask = work["Lifecycle Type"].isin(
        ["New / Coming Soon", "Discontinued / Phase Out", "Display Low Stock", "Display No Stock", "Renewal"]
    )
    work.loc[override_mask, "Action"] = work.loc[override_mask, "Lifecycle Action"]
    return work


def build_purchase_summary(
    purchase_file: str | Path | BinaryIO | BytesIO | list[str | Path | BinaryIO | BytesIO] | None,
    purchase_filter_keyword: str = "",
    purchase_years: list[int] | None = None,
    location_filter: str = "",
) -> tuple[pd.DataFrame | None, pd.DataFrame | None, pd.DataFrame | None, pd.DataFrame | None, pd.DataFrame | None, pd.DataFrame | None, dict]:
    if purchase_file is None:
        return None, None, None, None, None, None, {}
    files = purchase_file if isinstance(purchase_file, list) else [purchase_file]
    frames = []
    for idx, file in enumerate(files):
        frame = read_excel_with_detected_header(file)
        if purchase_years and idx < len(purchase_years):
            frame["_source_year"] = purchase_years[idx]
        frames.append(frame)
    df = pd.concat(frames, ignore_index=True)
    date_col = detect_column(
        df,
        ["Creation date", "PO Date", "Purchase Date", "Order Date", "Invoice Date", "Date", "Created At", "Month, Year"],
        required=False,
    )
    year_col = detect_column(df, ["Year", "Purchase Year", "PO Year"], required=False)
    amount_col = detect_column(
        df,
        ["Total.2", "Total.1", "Line Total", "Line Amount", "Purchase Amount", "Amount", "Total Amount", "Net Amount", "Grand Total", "Subtotal", "Total without taxes", "Total"],
        required=False,
    )
    qty_col = detect_column(df, ["Quantity", "Qty", "Order Qty", "Purchase Qty", "Units"], required=False)
    unit_cost_col = detect_column(
        df,
        ["Unit Cost", "Unit Price", "Wholesale Price", "Cost", "Price", "FOB", "Purchase Price"],
        required=False,
    )
    sku_col = detect_column(df, ["SKU", "Product SKU", "Item SKU", "JAN", "Barcode", "Product Code"], required=False)
    product_col = detect_column(df, ["Product", "Product Name", "Product / Service", "Description", "Item Name"], required=False)
    barcode_col = detect_column(df, ["Barcode", "JAN", "UPC", "EAN"], required=False)
    location_col = detect_column(df, ["Location", "Warehouse", "Receiving Location", "Stock Location"], required=False)

    detected = {
        "purchase_date": date_col or "Not detected",
        "purchase_year": year_col or "Not detected",
        "purchase_amount": amount_col or "Not detected",
        "purchase_qty": qty_col or "Not detected",
        "purchase_unit_cost": unit_cost_col or "Not detected",
        "purchase_sku": sku_col or "Not detected",
        "purchase_product": product_col or "Not detected",
        "purchase_barcode": barcode_col or "Not detected",
        "purchase_location": location_col or "Not detected",
        "purchase_location_filter": location_filter or "Not used",
        "purchase_filter_keyword": purchase_filter_keyword or "Not used",
        "purchase_year_mode": "Manual upload year" if purchase_years else "Detected from PO date/year",
    }

    if amount_col is None and not (qty_col and unit_cost_col):
        raise ValueError("Purchase report needs an Amount column, or both Qty and Unit Cost columns.")
    if not purchase_years and date_col is None and year_col is None:
        raise ValueError("Purchase report needs a Date/PO Date column or a Year column.")

    work = df.copy()
    raw_records = len(work)
    keyword = purchase_filter_keyword.strip()
    if keyword:
        text_cols = [c for c in [product_col, sku_col, barcode_col, detect_column(df, ["Supplier", "Brand", "Vendor"], required=False)] if c]
        if not text_cols:
            raise ValueError("Purchase keyword filter was provided, but no Product/SKU/Barcode/Supplier text columns were detected.")
        mask = pd.Series(False, index=work.index)
        for col in text_cols:
            mask = mask | work[col].astype(str).str.contains(keyword, case=False, na=False)
        work = work[mask].copy()
    if location_filter.strip():
        if not location_col:
            raise ValueError("Location filter was provided, but no Location/Warehouse column was detected in the Purchase report.")
        work = filter_by_location(work, location_col, location_filter)
    if purchase_years and "_source_year" in work.columns:
        work["_year"] = pd.to_numeric(work["_source_year"], errors="coerce")
        work["_date"] = pd.to_datetime(work[date_col], errors="coerce") if date_col else pd.NaT
    elif date_col:
        work["_date"] = pd.to_datetime(work[date_col], errors="coerce")
        work["_year"] = work["_date"].dt.year
    else:
        work["_date"] = pd.NaT
        work["_year"] = pd.to_numeric(work[year_col], errors="coerce")

    if amount_col:
        work["_purchase_amount"] = pd.to_numeric(work[amount_col], errors="coerce").fillna(0)
    else:
        work["_purchase_amount"] = (
            pd.to_numeric(work[qty_col], errors="coerce").fillna(0)
            * pd.to_numeric(work[unit_cost_col], errors="coerce").fillna(0)
        )

    current_year = datetime.now().year
    current_date = pd.Timestamp(datetime.now().date())
    rows = []
    for year in [2024, 2025, 2026]:
        subset = work[work["_year"] == year].copy()
        if year == current_year and date_col:
            subset = subset[subset["_date"].isna() | (subset["_date"] <= current_date)]
            period = f"{year} YTD through {current_date.date()}"
        elif year == current_year:
            period = f"{year} YTD"
        else:
            period = f"{year} full year"
        rows.append(
            {
                "Year": year,
                "Period": period,
                "Purchase Amount": float(subset["_purchase_amount"].sum()),
                "Record Count": int(len(subset)),
            }
        )
    summary = pd.DataFrame(rows)
    sku_summary = None
    purchase_key_summary = None
    location_summary = None
    year_location_summary = None
    sku_location_summary = None
    if sku_col or product_col:
        group_cols = [c for c in [sku_col, product_col] if c]
        aggregations = {"Purchase Amount": ("_purchase_amount", "sum"), "Record Count": ("_purchase_amount", "count")}
        if qty_col:
            aggregations["Quantity"] = (qty_col, lambda s: pd.to_numeric(s, errors="coerce").fillna(0).sum())
        sku_summary = work.groupby(group_cols, dropna=False).agg(**aggregations).reset_index().sort_values("Purchase Amount", ascending=False)
        year_amount = work.pivot_table(index=group_cols, columns="_year", values="_purchase_amount", aggfunc="sum", fill_value=0).reset_index()
        year_amount.columns = [f"{int(col)} Purchase Amount" if isinstance(col, (int, float, np.integer, np.floating)) else col for col in year_amount.columns]
        if qty_col:
            work["_purchase_qty"] = pd.to_numeric(work[qty_col], errors="coerce").fillna(0)
            year_qty = work.pivot_table(index=group_cols, columns="_year", values="_purchase_qty", aggfunc="sum", fill_value=0).reset_index()
            year_qty.columns = [f"{int(col)} Purchase Qty" if isinstance(col, (int, float, np.integer, np.floating)) else col for col in year_qty.columns]
            sku_summary = sku_summary.merge(year_amount, on=group_cols, how="left").merge(year_qty, on=group_cols, how="left")
        else:
            sku_summary = sku_summary.merge(year_amount, on=group_cols, how="left")
        for year in [2024, 2025, 2026]:
            amount_col_name = f"{year} Purchase Amount"
            if amount_col_name not in sku_summary.columns:
                sku_summary[amount_col_name] = 0
            qty_col_name = f"{year} Purchase Qty"
            if qty_col and qty_col_name not in sku_summary.columns:
                sku_summary[qty_col_name] = 0
        ordered_cols = group_cols + ["Purchase Amount", "Record Count"]
        if qty_col and "Quantity" in sku_summary.columns:
            ordered_cols.append("Quantity")
        ordered_cols += [f"{year} Purchase Amount" for year in [2024, 2025, 2026]]
        ordered_cols += [f"{year} Purchase Qty" for year in [2024, 2025, 2026] if f"{year} Purchase Qty" in sku_summary.columns]
        sku_summary = sku_summary[[col for col in ordered_cols if col in sku_summary.columns]]

    if location_col:
        work["_location"] = work[location_col].map(clean_location)
        location_summary = (
            work.groupby("_location", dropna=False)
            .agg(Purchase_Amount=("_purchase_amount", "sum"), Record_Count=("_purchase_amount", "count"))
            .reset_index()
            .rename(columns={"_location": "Location", "Purchase_Amount": "Purchase Amount", "Record_Count": "Record Count"})
            .sort_values("Purchase Amount", ascending=False)
        )
        location_year = work.pivot_table(index="_location", columns="_year", values="_purchase_amount", aggfunc="sum", fill_value=0).reset_index()
        location_year.columns = [f"{int(col)} Purchase Amount" if isinstance(col, (int, float, np.integer, np.floating)) else "Location" for col in location_year.columns]
        location_summary = location_summary.merge(location_year, on="Location", how="left")
        for year in [2024, 2025, 2026]:
            col = f"{year} Purchase Amount"
            if col not in location_summary.columns:
                location_summary[col] = 0
        location_summary = location_summary[["Location", "Purchase Amount", "Record Count", "2024 Purchase Amount", "2025 Purchase Amount", "2026 Purchase Amount"]]

        aggregations = {"Purchase Amount": ("_purchase_amount", "sum"), "Record Count": ("_purchase_amount", "count")}
        if qty_col:
            work["_purchase_qty"] = pd.to_numeric(work[qty_col], errors="coerce").fillna(0)
            aggregations["Purchase Qty"] = ("_purchase_qty", "sum")
        year_location_summary = (
            work[work["_year"].notna()]
            .groupby(["_year", "_location"], dropna=False)
            .agg(**aggregations)
            .reset_index()
            .rename(columns={"_year": "Year", "_location": "Location"})
            .sort_values(["Year", "Purchase Amount"], ascending=[False, False])
        )
        if year_location_summary is not None and not year_location_summary.empty:
            year_location_summary["Year"] = year_location_summary["Year"].astype(int)

        if sku_col or product_col:
            sku_loc_group_cols = [c for c in [sku_col, product_col, "_location"] if c]
            sku_location_summary = (
                work.groupby(sku_loc_group_cols, dropna=False)
                .agg(Purchase_Amount=("_purchase_amount", "sum"), Record_Count=("_purchase_amount", "count"))
                .reset_index()
                .rename(columns={"_location": "Location", "Purchase_Amount": "Purchase Amount", "Record_Count": "Record Count"})
                .sort_values("Purchase Amount", ascending=False)
            )

    key_records = []
    key_cols = [c for c in [sku_col, barcode_col] if c]
    work["_line_id"] = range(len(work))
    for key_col in key_cols:
        temp = work[["_line_id", key_col, "_year", "_purchase_amount"]].copy()
        temp["_purchase_key"] = temp[key_col].map(normalize_sku)
        temp = temp[temp["_purchase_key"] != ""]
        key_records.append(temp[["_line_id", "_purchase_key", "_year", "_purchase_amount"]])
    if key_records:
        key_work = pd.concat(key_records, ignore_index=True).drop_duplicates(subset=["_line_id", "_purchase_key"])
        purchase_key_summary = key_work.pivot_table(index="_purchase_key", columns="_year", values="_purchase_amount", aggfunc="sum", fill_value=0).reset_index()
        purchase_key_summary.columns = [
            f"{int(col)} Purchase Amount" if isinstance(col, (int, float, np.integer, np.floating)) else col for col in purchase_key_summary.columns
        ]
        for year in [2024, 2025, 2026]:
            col = f"{year} Purchase Amount"
            if col not in purchase_key_summary.columns:
                purchase_key_summary[col] = 0
        purchase_key_summary["Total Purchase Amount"] = purchase_key_summary[[f"{year} Purchase Amount" for year in [2024, 2025, 2026]]].sum(axis=1)
    detected["purchase_total_records_before_filter"] = raw_records
    detected["purchase_total_records_after_filter"] = len(work)
    return summary, sku_summary, purchase_key_summary, location_summary, year_location_summary, sku_location_summary, detected


def build_location_year_business_view(
    sales_year_location_summary: pd.DataFrame | None,
    stock_location_summary: pd.DataFrame | None,
    purchase_year_location_summary: pd.DataFrame | None,
) -> pd.DataFrame | None:
    view = None
    if sales_year_location_summary is not None and not sales_year_location_summary.empty:
        sales_cols = ["Year", "Location", "Sales Qty"]
        if "Sales Amount" in sales_year_location_summary.columns:
            sales_cols.append("Sales Amount")
        view = sales_year_location_summary[sales_cols].copy()
    if purchase_year_location_summary is not None and not purchase_year_location_summary.empty:
        po_cols = ["Year", "Location", "Purchase Amount", "Record Count"]
        if "Purchase Qty" in purchase_year_location_summary.columns:
            po_cols.append("Purchase Qty")
        purchase_view = purchase_year_location_summary[po_cols].copy()
        view = purchase_view if view is None else view.merge(purchase_view, on=["Year", "Location"], how="outer")
    if view is None or view.empty:
        return None
    if stock_location_summary is not None and not stock_location_summary.empty:
        stock_cols = ["Location", "Available", "Incoming", "On Hand", "Future Inventory"]
        stock_cols = [col for col in stock_cols if col in stock_location_summary.columns]
        view = view.merge(stock_location_summary[stock_cols], on="Location", how="left")
    for col in ["Sales Qty", "Sales Amount", "Purchase Amount", "Purchase Qty", "Available", "Incoming", "On Hand", "Future Inventory"]:
        if col in view.columns:
            view[col] = pd.to_numeric(view[col], errors="coerce").fillna(0)
    ordered = [
        "Year",
        "Location",
        "Sales Qty",
        "Sales Amount",
        "Purchase Amount",
        "Purchase Qty",
        "Available",
        "Incoming",
        "On Hand",
        "Future Inventory",
        "Record Count",
    ]
    ordered = [col for col in ordered if col in view.columns]
    return view[ordered].sort_values(["Year", "Location"], ascending=[False, True])


def build_sales_purchase_year_summary(
    sales_year_summary: pd.DataFrame | None,
    purchase_summary: pd.DataFrame | None,
) -> pd.DataFrame | None:
    if sales_year_summary is None and purchase_summary is None:
        return None
    view = None
    if sales_year_summary is not None and not sales_year_summary.empty:
        cols = ["Year", "SKU Count", "Sales Qty"]
        if "Sales Amount" in sales_year_summary.columns:
            cols.append("Sales Amount")
        view = sales_year_summary[cols].copy()
    if purchase_summary is not None and not purchase_summary.empty:
        purchase_view = purchase_summary[["Year", "Purchase Amount", "Record Count"]].copy()
        view = purchase_view if view is None else view.merge(purchase_view, on="Year", how="outer")
    if view is None or view.empty:
        return None
    for col in ["SKU Count", "Sales Qty", "Sales Amount", "Purchase Amount", "Record Count"]:
        if col in view.columns:
            view[col] = pd.to_numeric(view[col], errors="coerce").fillna(0)
    if {"Purchase Amount", "Sales Qty"}.issubset(view.columns):
        view["Purchase Amount / Sales Qty"] = np.where(view["Sales Qty"] > 0, view["Purchase Amount"] / view["Sales Qty"], np.nan)
    if {"Purchase Amount", "Sales Amount"}.issubset(view.columns):
        view["Purchase / Sales Amount"] = np.where(view["Sales Amount"] > 0, view["Purchase Amount"] / view["Sales Amount"], np.nan)
    ordered = [
        "Year",
        "SKU Count",
        "Sales Qty",
        "Sales Amount",
        "Purchase Amount",
        "Purchase Amount / Sales Qty",
        "Purchase / Sales Amount",
        "Record Count",
    ]
    return view[[col for col in ordered if col in view.columns]].sort_values("Year")


def build_analysis(
    sales_file: str | Path | BinaryIO | BytesIO | list[str | Path | BinaryIO | BytesIO],
    stock_file: str | Path | BinaryIO | BytesIO,
    catalogue_file: str | Path | BinaryIO | BytesIO | None = None,
    purchase_file: str | Path | BinaryIO | BytesIO | list[str | Path | BinaryIO | BytesIO] | None = None,
    purchase_filter_keyword: str = "",
    purchase_years: list[int] | None = None,
    sales_years: list[int] | None = None,
    location_filter: str = "",
) -> AnalysisResult:
    sales_df = read_year_labeled_files(sales_file, sales_years) if isinstance(sales_file, list) else read_excel_any(sales_file)
    stock_df = read_excel_any(stock_file)
    catalogue_df = read_excel_with_detected_header(catalogue_file) if catalogue_file else None
    purchase_summary, purchase_sku_summary, purchase_key_summary, purchase_location_summary, purchase_year_location_summary, purchase_sku_location_summary, purchase_cols = build_purchase_summary(
        purchase_file, purchase_filter_keyword, purchase_years, location_filter
    )

    sales, sales_year_summary, sales_location_summary, sales_year_location_summary, sales_cols = aggregate_sales(
        sales_df,
        purchase_filter_keyword,
        manual_year_mode=bool(sales_years),
        location_filter=location_filter,
    )
    stock, stock_location_summary, stock_cols = aggregate_stock(stock_df, purchase_filter_keyword, location_filter)
    sales_purchase_year_summary = build_sales_purchase_year_summary(sales_year_summary, purchase_summary)
    location_year_business_view = build_location_year_business_view(
        sales_year_location_summary,
        stock_location_summary,
        purchase_year_location_summary,
    )

    final = sales.merge(stock, on="Product SKU", how="outer")
    final["Qty"] = pd.to_numeric(final["Qty"], errors="coerce").fillna(0)
    final["Product Name"] = final["Product Name"].fillna(final.get("Stock Product Name", "")).fillna("")
    final = final.drop(columns=["Stock Product Name"], errors="ignore")
    for col in ["Available", "Incoming", "On Hand"]:
        final[col] = pd.to_numeric(final.get(col, 0), errors="coerce").fillna(0)

    final = final.sort_values("Qty", ascending=False).reset_index(drop=True)
    total_qty = final["Qty"].sum()
    final["Contribution %"] = np.where(total_qty > 0, final["Qty"] / total_qty, 0)
    final["Cumulative %"] = final["Contribution %"].cumsum()
    final["SABC Type"] = np.select(
        [
            final["Cumulative %"] <= 0.05,
            final["Cumulative %"] <= 0.80,
            final["Cumulative %"] <= 0.95,
        ],
        ["S", "A", "B"],
        default="C",
    )
    final.loc[final["Qty"] <= 0, "SABC Type"] = "C"

    final["Future Inventory"] = final["Available"] + final["Incoming"]
    final["Adjusted Future Inventory"] = final["Future Inventory"].clip(lower=0)
    final["Avg Monthly Sales"] = final["Qty"] / 12
    final["Coverage"] = np.where(final["Avg Monthly Sales"] > 0, final["Adjusted Future Inventory"] / final["Avg Monthly Sales"], np.nan)

    conditions = [
        final["Avg Monthly Sales"] <= 0,
        (final["Adjusted Future Inventory"] <= 0) & (final["Avg Monthly Sales"] > 0),
        final["Coverage"] < 1,
        final["Coverage"] < 3,
        final["Coverage"] < 6,
        final["Coverage"] >= 6,
    ]
    final["Inventory Status"] = np.select(
        conditions,
        ["No Sales / Review", "Stockout", "Urgent", "Healthy", "Monitor", "Overstock"],
        default="No Sales / Review",
    )
    action_map = {
        "Stockout": "Immediate Replenishment",
        "Urgent": "Replenish",
        "Healthy": "Monitor",
        "Monitor": "Review PO",
        "Overstock": "Reduce PO",
        "No Sales / Review": "Review SKU",
    }
    final["Action"] = final["Inventory Status"].map(action_map)
    final, catalogue_cols = apply_catalogue(final, catalogue_df)
    final = apply_lifecycle_strategy(final)
    if purchase_key_summary is not None:
        final = final.merge(purchase_key_summary, left_on="Product SKU", right_on="_purchase_key", how="left").drop(columns=["_purchase_key"], errors="ignore")
        for col in ["2024 Purchase Amount", "2025 Purchase Amount", "2026 Purchase Amount", "Total Purchase Amount"]:
            final[col] = pd.to_numeric(final.get(col, 0), errors="coerce").fillna(0)

    base_cols = [
        "Product SKU",
        "Product Name",
        "Qty",
        "Contribution %",
        "Cumulative %",
        "SABC Type",
        "Available",
        "Incoming",
        "Future Inventory",
        "Adjusted Future Inventory",
        "Avg Monthly Sales",
        "Coverage",
        "Inventory Status",
        "Action",
    ]
    purchase_cols_in_final = [c for c in ["2024 Purchase Amount", "2025 Purchase Amount", "2026 Purchase Amount", "Total Purchase Amount"] if c in final.columns]
    base_cols = base_cols + purchase_cols_in_final
    lifecycle_cols = [c for c in ["Recommendation / Lifecycle Status", "Lifecycle Type", "Lifecycle Action", "Lifecycle Note"] if c in final.columns]
    base_cols = base_cols + lifecycle_cols
    extra_cols = [c for c in final.columns if c not in base_cols + ["On Hand"]]
    final = final[base_cols + extra_cols]

    sabc_summary = summarize(final, "SABC Type")
    inventory_status_summary = summarize(final, "Inventory Status")
    action_summary = summarize(final, "Action")
    matrix = pd.crosstab(final["SABC Type"], final["Inventory Status"], values=final["Product SKU"], aggfunc="count").fillna(0).astype(int)
    matrix = matrix.reindex(index=["S", "A", "B", "C"], fill_value=0)

    insights = build_insights(final, sabc_summary, inventory_status_summary, action_summary)
    if purchase_summary is not None:
        insights["purchase_summary"] = purchase_summary
        insights["purchase_sku_summary"] = purchase_sku_summary
        insights["purchase_location_summary"] = purchase_location_summary
        insights["purchase_year_location_summary"] = purchase_year_location_summary
        insights["purchase_sku_location_summary"] = purchase_sku_location_summary
        insights["purchase_total"] = float(purchase_summary["Purchase Amount"].sum())
    if sales_location_summary is not None:
        insights["sales_location_summary"] = sales_location_summary
    if sales_year_summary is not None:
        insights["sales_year_summary"] = sales_year_summary
    if sales_purchase_year_summary is not None:
        insights["sales_purchase_year_summary"] = sales_purchase_year_summary
    if sales_year_location_summary is not None:
        insights["sales_year_location_summary"] = sales_year_location_summary
    if stock_location_summary is not None:
        insights["stock_location_summary"] = stock_location_summary
    if location_year_business_view is not None:
        insights["location_year_business_view"] = location_year_business_view
    return AnalysisResult(
        final,
        sabc_summary,
        inventory_status_summary,
        action_summary,
        matrix,
        sales_year_summary,
        sales_purchase_year_summary,
        sales_location_summary,
        sales_year_location_summary,
        stock_location_summary,
        purchase_summary,
        purchase_sku_summary,
        purchase_location_summary,
        purchase_year_location_summary,
        purchase_sku_location_summary,
        location_year_business_view,
        insights,
        {**sales_cols, **stock_cols, **catalogue_cols, **purchase_cols},
    )


def summarize(final: pd.DataFrame, group_col: str) -> pd.DataFrame:
    out = (
        final.groupby(group_col, dropna=False)
        .agg(
            SKU_Count=("Product SKU", "count"),
            Qty=("Qty", "sum"),
            Avg_Coverage=("Coverage", "mean"),
            Future_Inventory=("Adjusted Future Inventory", "sum"),
        )
        .reset_index()
    )
    total_qty = final["Qty"].sum()
    out["Qty Share"] = np.where(total_qty > 0, out["Qty"] / total_qty, 0)
    return out


def build_insights(final: pd.DataFrame, sabc: pd.DataFrame, status: pd.DataFrame, action: pd.DataFrame) -> dict:
    total_skus = len(final)
    total_qty = float(final["Qty"].sum())
    top_sku = final.iloc[0] if total_skus else None
    urgent = final[final["Inventory Status"].isin(["Stockout", "Urgent"])].sort_values(["SABC Type", "Qty"], ascending=[True, False])
    overstock = final[final["Inventory Status"] == "Overstock"].sort_values("Adjusted Future Inventory", ascending=False)
    core = final[final["SABC Type"].isin(["S", "A"])]
    no_sales = final[final["Inventory Status"] == "No Sales / Review"]
    top_sku_table = final.sort_values("Qty", ascending=False).head(5).copy()
    catalogue_present = any(col in final.columns for col in ["Category", "Recommendation / Lifecycle Status", "RP USD", "Wholesale Price", "Inner Case", "Outer Case"])
    catalogue_insights = build_catalogue_insights(final) if catalogue_present else {}
    return {
        "total_skus": total_skus,
        "total_qty": total_qty,
        "top_sku": "" if top_sku is None else top_sku["Product SKU"],
        "top_product": "" if top_sku is None else top_sku["Product Name"],
        "top_qty": 0 if top_sku is None else float(top_sku["Qty"]),
        "top_share": 0 if total_qty == 0 or top_sku is None else float(top_sku["Qty"] / total_qty),
        "urgent_count": len(urgent),
        "overstock_count": len(overstock),
        "no_sales_count": len(no_sales),
        "core_sku_count": len(core),
        "core_qty_share": 0 if total_qty == 0 else float(core["Qty"].sum() / total_qty),
        "top_sku_table": top_sku_table,
        "urgent_table": urgent.head(10),
        "overstock_table": overstock.head(10),
        "catalogue_present": catalogue_present,
        **catalogue_insights,
    }


def build_catalogue_insights(final: pd.DataFrame) -> dict:
    insights: dict = {}
    work = final.copy()

    if "Category" in work.columns:
        category = (
            work.groupby("Category", dropna=False)
            .agg(
                SKU_Count=("Product SKU", "count"),
                Qty=("Qty", "sum"),
                Avg_Coverage=("Coverage", "mean"),
                Future_Inventory=("Adjusted Future Inventory", "sum"),
            )
            .reset_index()
            .sort_values("Qty", ascending=False)
        )
        category["Qty Share"] = np.where(work["Qty"].sum() > 0, category["Qty"] / work["Qty"].sum(), 0)
        insights["category_summary"] = category.head(10)

    if "Recommendation / Lifecycle Status" in work.columns:
        lifecycle = (
            work.groupby(["Lifecycle Type", "Recommendation / Lifecycle Status"], dropna=False)
            .agg(SKU_Count=("Product SKU", "count"), Qty=("Qty", "sum"), Avg_Coverage=("Coverage", "mean"))
            .reset_index()
            .sort_values("Qty", ascending=False)
        )
        insights["lifecycle_summary"] = lifecycle.head(10)

    if "Lifecycle Type" in work.columns:
        lifecycle_cols = [
            "Product SKU",
            "Product Name",
            "Qty",
            "SABC Type",
            "Inventory Status",
            "Action",
            "Recommendation / Lifecycle Status",
            "Lifecycle Type",
            "Lifecycle Note",
        ]
        lifecycle_cols = [c for c in lifecycle_cols if c in work.columns]
        new_skus = work[work["Lifecycle Type"] == "New / Coming Soon"].copy()
        discontinued = work[(work["Lifecycle Type"] == "Discontinued / Phase Out") & (work["Qty"] > 0)].copy()
        display_risk = work[work["Lifecycle Type"].isin(["Display Low Stock", "Display No Stock"])].copy()
        renewal = work[work["Lifecycle Type"] == "Renewal"].copy()
        insights["new_sku_count"] = len(new_skus)
        insights["discontinued_with_sales_count"] = len(discontinued)
        insights["display_risk_count"] = len(display_risk)
        insights["renewal_count"] = len(renewal)
        insights["new_sku_table"] = new_skus.sort_values(["SABC Type", "Product SKU"])[lifecycle_cols].head(10)
        insights["discontinued_with_sales_table"] = discontinued.sort_values("Qty", ascending=False)[lifecycle_cols].head(10)
        insights["display_risk_table"] = display_risk.sort_values(["SABC Type", "Qty"], ascending=[True, False])[lifecycle_cols].head(10)
        insights["renewal_table"] = renewal.sort_values(["SABC Type", "Qty"], ascending=[True, False])[lifecycle_cols].head(10)

    if {"RP USD", "Wholesale Price"}.issubset(work.columns):
        work["RP USD"] = pd.to_numeric(work["RP USD"], errors="coerce")
        work["Wholesale Price"] = pd.to_numeric(work["Wholesale Price"], errors="coerce")
        if "Margin %" not in work.columns:
            work["Margin %"] = np.where(work["RP USD"] > 0, (work["RP USD"] - work["Wholesale Price"]) / work["RP USD"], np.nan)
        work["Estimated Wholesale Sales"] = work["Qty"] * work["Wholesale Price"]
        work["Estimated Retail Sales"] = work["Qty"] * work["RP USD"]
        priced = work.dropna(subset=["RP USD", "Wholesale Price"])
        insights["priced_sku_count"] = len(priced)
        insights["avg_margin"] = float(priced["Margin %"].mean()) if not priced.empty and "Margin %" in priced.columns else np.nan
        insights["estimated_wholesale_sales"] = float(priced["Estimated Wholesale Sales"].sum()) if not priced.empty else 0
        insights["estimated_retail_sales"] = float(priced["Estimated Retail Sales"].sum()) if not priced.empty else 0
        margin_cols = ["Product SKU", "Product Name", "Qty", "RP USD", "Wholesale Price", "Margin %", "SABC Type", "Inventory Status", "Action"]
        insights["low_margin_table"] = priced.sort_values("Margin %", ascending=True)[margin_cols].head(10)
        insights["top_value_table"] = priced.sort_values("Estimated Wholesale Sales", ascending=False)[
            ["Product SKU", "Product Name", "Qty", "Wholesale Price", "Estimated Wholesale Sales", "Margin %", "SABC Type", "Inventory Status"]
        ].head(10)

    if "Wholesale Price JPY" in work.columns:
        work["Wholesale Price JPY"] = pd.to_numeric(work["Wholesale Price JPY"], errors="coerce")
        priced_jpy = work.dropna(subset=["Wholesale Price JPY"]).copy()
        priced_jpy["Estimated Wholesale Sales JPY"] = priced_jpy["Qty"] * priced_jpy["Wholesale Price JPY"]
        insights["priced_jpy_sku_count"] = len(priced_jpy)
        insights["estimated_wholesale_sales_jpy"] = float(priced_jpy["Estimated Wholesale Sales JPY"].sum()) if not priced_jpy.empty else 0
        insights["top_value_jpy_table"] = priced_jpy.sort_values("Estimated Wholesale Sales JPY", ascending=False)[
            ["Product SKU", "Product Name", "Qty", "Wholesale Price JPY", "Estimated Wholesale Sales JPY", "SABC Type", "Inventory Status", "Action"]
        ].head(10)

    if "MOQ risk" in work.columns:
        moq = work[work["MOQ risk"].astype(str).str.len() > 0].copy()
        insights["moq_risk_count"] = len(moq)
        moq_cols = [c for c in ["Product SKU", "Product Name", "Qty", "Avg Monthly Sales", "Outer Case", "Coverage", "SABC Type", "MOQ risk"] if c in moq.columns]
        insights["moq_risk_table"] = moq.sort_values(["SABC Type", "Avg Monthly Sales"], ascending=[True, False])[moq_cols].head(10)
    else:
        insights["moq_risk_count"] = 0

    return insights


def dataframe_to_rows(df: pd.DataFrame) -> list[list[object]]:
    clean = df.copy()
    clean = clean.replace([np.inf, -np.inf], np.nan)
    clean = clean.where(pd.notna(clean), None)
    return [list(clean.columns)] + clean.values.tolist()


def write_df(ws, df: pd.DataFrame, table_name: str, percent_cols: Iterable[str] = ()):
    rows = dataframe_to_rows(df)
    for row in rows:
        ws.append(row)
    max_row, max_col = ws.max_row, ws.max_column
    ref = f"A1:{get_column_letter(max_col)}{max_row}"
    table = Table(displayName=table_name, ref=ref)
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True, showColumnStripes=False)
    ws.add_table(table)
    style_sheet(ws, percent_cols)


def style_sheet(ws, percent_cols: Iterable[str] = ()):
    header_fill = PatternFill("solid", fgColor=ACCENT)
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D2D2D7")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows():
        for cell in row:
            cell.border = Border(bottom=thin)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    percent_set = set(percent_cols)
    for col_idx, col_cells in enumerate(ws.iter_cols(min_row=1, max_row=ws.max_row), start=1):
        header = col_cells[0].value
        width = min(max(len(str(header or "")) + 2, 12), 32)
        for cell in col_cells[1:]:
            if isinstance(cell.value, str):
                width = min(max(width, len(cell.value[:40]) + 2), 42)
            if header in percent_set:
                cell.number_format = "0.0%"
            elif isinstance(cell.value, (int, float)):
                cell.number_format = "#,##0.00" if header in {"Coverage", "Avg Monthly Sales", "Avg_Coverage"} else "#,##0"
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def add_excel_dashboard_charts(wb: Workbook):
    ws = wb["Final Analysis"]
    status_ws = wb["Inventory Status Summary"]
    sabc_ws = wb["SABC Summary"]
    if status_ws.max_row > 1:
        chart = BarChart()
        chart.title = "Inventory Status by SKU Count"
        chart.y_axis.title = "SKU Count"
        chart.x_axis.title = "Status"
        data = Reference(status_ws, min_col=2, min_row=1, max_row=status_ws.max_row)
        cats = Reference(status_ws, min_col=1, min_row=2, max_row=status_ws.max_row)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(cats)
        chart.height = 7
        chart.width = 12
        ws.add_chart(chart, "P2")
    if sabc_ws.max_row > 1:
        pie = PieChart()
        pie.title = "Qty Share by SABC"
        data = Reference(sabc_ws, min_col=6, min_row=1, max_row=sabc_ws.max_row)
        cats = Reference(sabc_ws, min_col=1, min_row=2, max_row=sabc_ws.max_row)
        pie.add_data(data, titles_from_data=True)
        pie.set_categories(cats)
        pie.height = 7
        pie.width = 10
        ws.add_chart(pie, "P18")


def generate_excel(result: AnalysisResult) -> bytes:
    wb = Workbook()
    wb.remove(wb.active)
    sheets = [
        ("Final Analysis", result.final, "FinalAnalysisTable", ["Contribution %", "Cumulative %", "Margin %"]),
        ("SABC Summary", result.sabc_summary, "SABCSummaryTable", ["Qty Share"]),
        ("Inventory Status Summary", result.inventory_status_summary, "InventoryStatusSummaryTable", ["Qty Share"]),
        ("Action Summary", result.action_summary, "ActionSummaryTable", ["Qty Share"]),
        ("SABC x Inventory Status Matrix", result.matrix.reset_index(), "SABCMatrixTable", []),
    ]
    if result.location_year_business_view is not None and not result.location_year_business_view.empty:
        sheets.append(("Year Location View", result.location_year_business_view, "YearLocationViewTable", []))
    if result.sales_purchase_year_summary is not None and not result.sales_purchase_year_summary.empty:
        sheets.append(("Sales PO Year Compare", result.sales_purchase_year_summary, "SalesPOYearCompareTable", []))
    if result.sales_year_summary is not None and not result.sales_year_summary.empty:
        sheets.append(("Sales Summary by Year", result.sales_year_summary, "SalesSummaryByYearTable", []))
    if result.sales_location_summary is not None and not result.sales_location_summary.empty:
        sheets.append(("Sales by Location", result.sales_location_summary, "SalesByLocationTable", []))
    if result.sales_year_location_summary is not None and not result.sales_year_location_summary.empty:
        sheets.append(("Sales Year Location", result.sales_year_location_summary, "SalesYearLocationTable", []))
    if result.stock_location_summary is not None and not result.stock_location_summary.empty:
        sheets.append(("Stock by Location", result.stock_location_summary, "StockByLocationTable", []))
    if result.purchase_summary is not None:
        sheets.append(("Purchase Summary", result.purchase_summary, "PurchaseSummaryTable", []))
        if result.purchase_sku_summary is not None and not result.purchase_sku_summary.empty:
            sheets.append(("Purchase by SKU", result.purchase_sku_summary, "PurchaseBySKUTable", []))
        if result.purchase_location_summary is not None and not result.purchase_location_summary.empty:
            sheets.append(("Purchase by Location", result.purchase_location_summary, "PurchaseByLocationTable", []))
        if result.purchase_year_location_summary is not None and not result.purchase_year_location_summary.empty:
            sheets.append(("Purchase Year Location", result.purchase_year_location_summary, "PurchaseYearLocationTable", []))
        if result.purchase_sku_location_summary is not None and not result.purchase_sku_location_summary.empty:
            sheets.append(("Purchase SKU Location", result.purchase_sku_location_summary, "PurchaseSKULocationTable", []))
    if result.insights.get("catalogue_present"):
        if "category_summary" in result.insights:
            sheets.append(("Category Summary", result.insights["category_summary"], "CategorySummaryTable", ["Qty Share"]))
        if "lifecycle_summary" in result.insights:
            sheets.append(("Lifecycle Summary", result.insights["lifecycle_summary"], "LifecycleSummaryTable", []))
        if "new_sku_table" in result.insights and not result.insights["new_sku_table"].empty:
            sheets.append(("New Coming Soon", result.insights["new_sku_table"], "NewComingSoonTable", []))
        if "discontinued_with_sales_table" in result.insights and not result.insights["discontinued_with_sales_table"].empty:
            sheets.append(("Discontinued Sales", result.insights["discontinued_with_sales_table"], "DiscontinuedSalesTable", []))
        if "display_risk_table" in result.insights and not result.insights["display_risk_table"].empty:
            sheets.append(("Display Risk", result.insights["display_risk_table"], "DisplayRiskTable", []))
        if "renewal_table" in result.insights and not result.insights["renewal_table"].empty:
            sheets.append(("Renewal Transition", result.insights["renewal_table"], "RenewalTransitionTable", []))
        if "top_value_table" in result.insights:
            sheets.append(("Top Value SKUs", result.insights["top_value_table"], "TopValueTable", ["Margin %"]))
        if "top_value_jpy_table" in result.insights:
            sheets.append(("Top Value JPY", result.insights["top_value_jpy_table"], "TopValueJPYTable", []))
        if "low_margin_table" in result.insights:
            sheets.append(("Low Margin Review", result.insights["low_margin_table"], "LowMarginTable", ["Margin %"]))
        if "moq_risk_table" in result.insights and not result.insights["moq_risk_table"].empty:
            sheets.append(("MOQ Risk", result.insights["moq_risk_table"], "MOQRiskTable", []))
    for name, df, table_name, percent_cols in sheets:
        ws = wb.create_sheet(name)
        write_df(ws, df, table_name, percent_cols)

    ws = wb["Final Analysis"]
    headers = [cell.value for cell in ws[1]]
    if "Inventory Status" in headers:
        col = get_column_letter(headers.index("Inventory Status") + 1)
        rng = f"{col}2:{col}{ws.max_row}"
        ws.conditional_formatting.add(rng, CellIsRule(operator="equal", formula=['"Stockout"'], fill=PatternFill("solid", fgColor="FCE4E4")))
        ws.conditional_formatting.add(rng, CellIsRule(operator="equal", formula=['"Urgent"'], fill=PatternFill("solid", fgColor="FFF4CC")))
        ws.conditional_formatting.add(rng, CellIsRule(operator="equal", formula=['"Overstock"'], fill=PatternFill("solid", fgColor="EAF4FF")))
    add_excel_dashboard_charts(wb)
    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


def try_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def make_bar_chart_png(df: pd.DataFrame, label_col: str, value_col: str, title: str) -> str:
    data = df[[label_col, value_col]].copy().head(8)
    data[value_col] = pd.to_numeric(data[value_col], errors="coerce").fillna(0)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    width, height = 900, 420
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = try_font(24, True)
    label_font = try_font(15)
    small_font = try_font(13)
    draw.text((32, 24), title, fill=f"#{ACCENT}", font=title_font)
    max_value = max(float(data[value_col].max()), 1.0)
    top, left, bar_h, gap = 82, 240, 30, 16
    for i, (_, row) in enumerate(data.iterrows()):
        y = top + i * (bar_h + gap)
        label = str(row[label_col])[:26]
        value = float(row[value_col])
        draw.text((32, y + 5), label, fill=f"#{ACCENT}", font=label_font)
        draw.rounded_rectangle((left, y, width - 90, y + bar_h), radius=8, fill=f"#{GRAY}")
        bar_w = int((width - 90 - left) * value / max_value)
        draw.rounded_rectangle((left, y, left + bar_w, y + bar_h), radius=8, fill=f"#{BLUE}")
        draw.text((left + bar_w + 10, y + 6), f"{value:,.0f}", fill=f"#{ACCENT}", font=small_font)
    img.save(tmp.name, "PNG")
    return tmp.name


def make_sabc_chart_png(df: pd.DataFrame) -> str:
    return make_bar_chart_png(df.sort_values("Qty", ascending=False), "SABC Type", "Qty", "Sales Quantity by SABC Type")


def set_doc_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(29, 29, 31)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, df: pd.DataFrame, max_rows: int = 10):
    clean = df.head(max_rows).copy().replace([np.inf, -np.inf], np.nan)
    clean = clean.where(pd.notna(clean), "")
    table = doc.add_table(rows=1, cols=len(clean.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(clean.columns):
        hdr[i].text = str(col)
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr[i]._tc.get_or_add_tcPr().append(_cell_shading(ACCENT))
    for _, row in clean.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                text = f"{value:.1%}" if "Share" in str(clean.columns[i]) or "%" in str(clean.columns[i]) else f"{value:,.2f}"
            else:
                text = str(value)
            cells[i].text = text
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def _cell_shading(fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_key_value_paragraph(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)


def add_catalogue_report_sections(doc: Document, insights: dict):
    doc.add_heading("Catalogue-Enhanced Commercial Analysis", level=1)
    doc.add_paragraph(
        "Because a catalogue / price list was provided, this report adds a commercial layer on top of sales and inventory: "
        "category performance, lifecycle status, retail price, wholesale price, margin, case pack, and MOQ risk where those fields are available."
    )

    if "category_summary" in insights and not insights["category_summary"].empty:
        doc.add_heading("Category Performance", level=2)
        doc.add_paragraph("This view shows which product categories carry sales volume and whether their inventory coverage is balanced.")
        add_table(doc, insights["category_summary"], max_rows=10)

    if "lifecycle_summary" in insights and not insights["lifecycle_summary"].empty:
        doc.add_heading("Recommendation / Lifecycle Status", level=2)
        doc.add_paragraph("This view helps compare current lifecycle labels against real sales and coverage behavior.")
        add_table(doc, insights["lifecycle_summary"], max_rows=10)

    if any(insights.get(key, 0) for key in ["new_sku_count", "discontinued_with_sales_count", "display_risk_count", "renewal_count"]):
        doc.add_heading("Lifecycle / Recommendation Strategy", level=2)
        add_bullet(doc, f"New / Coming Soon SKUs: {insights.get('new_sku_count', 0):,}. These should be treated as launch planning items, not as failed no-sales SKUs.")
        add_bullet(doc, f"Discontinued / phase-out SKUs with sales history: {insights.get('discontinued_with_sales_count', 0):,}. These should not receive normal replenishment recommendations.")
        add_bullet(doc, f"Display low/no stock risks: {insights.get('display_risk_count', 0):,}. These may need tester, display, or shelf support from the manufacturer.")
        add_bullet(doc, f"Renewal SKUs: {insights.get('renewal_count', 0):,}. These need old/new version transition planning.")
        if "new_sku_table" in insights and not insights["new_sku_table"].empty:
            doc.add_heading("New / Coming Soon Launch Candidates", level=2)
            add_table(doc, insights["new_sku_table"], max_rows=10)
        if "discontinued_with_sales_table" in insights and not insights["discontinued_with_sales_table"].empty:
            doc.add_heading("Discontinued SKUs With Sales History", level=2)
            add_table(doc, insights["discontinued_with_sales_table"], max_rows=10)
        if "display_risk_table" in insights and not insights["display_risk_table"].empty:
            doc.add_heading("Display / Tester Support Risk", level=2)
            add_table(doc, insights["display_risk_table"], max_rows=10)
        if "renewal_table" in insights and not insights["renewal_table"].empty:
            doc.add_heading("Renewal Transition Plan", level=2)
            add_table(doc, insights["renewal_table"], max_rows=10)

    if "priced_sku_count" in insights:
        doc.add_heading("Pricing and Margin Analysis", level=2)
        add_bullet(doc, f"Priced SKUs matched: {insights.get('priced_sku_count', 0):,}.")
        add_bullet(doc, f"Estimated wholesale sales value from the uploaded sales Qty base: ${insights.get('estimated_wholesale_sales', 0):,.0f}.")
        add_bullet(doc, f"Estimated retail sales value from the uploaded sales Qty base: ${insights.get('estimated_retail_sales', 0):,.0f}.")
        avg_margin = insights.get("avg_margin", np.nan)
        if pd.notna(avg_margin):
            add_bullet(doc, f"Average gross margin based on RP USD and Wholesale Price: {avg_margin:.1%}.")
        if "top_value_table" in insights and not insights["top_value_table"].empty:
            doc.add_heading("Top Value SKUs", level=2)
            doc.add_paragraph("These SKUs are important not only by units sold, but also by estimated wholesale sales value.")
            add_table(doc, insights["top_value_table"], max_rows=10)
        if "low_margin_table" in insights and not insights["low_margin_table"].empty:
            doc.add_heading("Low-Margin Review", level=2)
            doc.add_paragraph("Low-margin SKUs should be reviewed before heavy promotion, bundle discounts, or large replenishment commitments.")
            add_table(doc, insights["low_margin_table"], max_rows=10)

    if "priced_jpy_sku_count" in insights:
        doc.add_heading("Wholesale Value Analysis (JPY)", level=2)
        add_bullet(doc, f"Wholesale JPY price matched for {insights.get('priced_jpy_sku_count', 0):,} SKUs.")
        add_bullet(doc, f"Estimated wholesale sales value from uploaded sales Qty: JPY {insights.get('estimated_wholesale_sales_jpy', 0):,.0f}.")
        add_bullet(doc, "Margin is not calculated from this file when RP is in USD and wholesale is in JPY, because mixed currencies would create a misleading margin.")
        if "top_value_jpy_table" in insights and not insights["top_value_jpy_table"].empty:
            doc.add_heading("Top Wholesale Value SKUs (JPY)", level=2)
            add_table(doc, insights["top_value_jpy_table"], max_rows=10)

    if insights.get("moq_risk_count", 0) > 0 and "moq_risk_table" in insights:
        doc.add_heading("MOQ / Case Pack Risk", level=2)
        doc.add_paragraph(
            "The table below flags SKUs where outer case quantity appears high versus monthly sales velocity. "
            "These SKUs may require smaller MOQ, split shipment, or conservative PO planning."
        )
        add_table(doc, insights["moq_risk_table"], max_rows=10)


def generate_word_report(result: AnalysisResult, brand_name: str = "Brand") -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"{brand_name} Business Analysis & 6-Month Plan")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(29, 29, 31)
    subtitle = doc.add_paragraph("Auto-generated from Sales Report and Stock Levels Report")
    subtitle.runs[0].font.color.rgb = RGBColor(134, 134, 139)

    insights = result.insights
    has_catalogue = bool(insights.get("catalogue_present"))
    doc.add_heading("Executive Summary", level=1)
    add_bullet(doc, f"Analyzed {insights['total_skus']:,} SKUs with {insights['total_qty']:,.0f} units sold in the uploaded sales period.")
    add_bullet(doc, f"Top SKU: {insights['top_sku']} - {insights['top_product']} ({insights['top_qty']:,.0f} units, {insights['top_share']:.1%} of sales).")
    add_bullet(
        doc,
        f"S/A core SKUs represent {insights['core_sku_count']:,} SKUs and {insights['core_qty_share']:.1%} of quantity sold. "
        "These items should be treated as the commercial priority group: protect availability, review PO timing first, and use them as the anchor products for channel growth."
    )
    add_bullet(doc, f"{insights['urgent_count']:,} SKUs need immediate replenishment attention; {insights['overstock_count']:,} SKUs show overstock risk.")
    if has_catalogue and "priced_sku_count" in insights:
        add_bullet(
            doc,
            f"Catalogue pricing was linked for {insights['priced_sku_count']:,} SKUs. Estimated wholesale sales value is "
            f"${insights.get('estimated_wholesale_sales', 0):,.0f}, with an average gross margin of {insights.get('avg_margin', 0):.1%} where price data is available."
        )
    if has_catalogue and insights.get("moq_risk_count", 0) > 0:
        add_bullet(doc, f"{insights['moq_risk_count']:,} SKUs show MOQ or outer-case risk, meaning case pack size may be too large versus current monthly sales velocity.")
    if has_catalogue and any(insights.get(key, 0) for key in ["new_sku_count", "discontinued_with_sales_count", "display_risk_count", "renewal_count"]):
        add_bullet(
            doc,
            f"Catalogue recommendation flags changed the business action for lifecycle-sensitive SKUs: "
            f"{insights.get('new_sku_count', 0):,} new/coming soon, "
            f"{insights.get('discontinued_with_sales_count', 0):,} discontinued with sales history, "
            f"{insights.get('display_risk_count', 0):,} display risk, and "
            f"{insights.get('renewal_count', 0):,} renewal items."
        )

    doc.add_heading("Top 5 SKUs by 12-Month Sales", level=2)
    top_cols = ["Product SKU", "Product Name", "Qty", "Contribution %", "Cumulative %", "SABC Type", "Coverage", "Inventory Status", "Action"]
    add_table(doc, insights["top_sku_table"][top_cols], max_rows=5)

    doc.add_heading("Data Sources", level=1)
    add_key_value_paragraph(doc, "Sales Report", "SKU-level quantity sold, product name, and month/year when available.")
    add_key_value_paragraph(doc, "Stock Levels Report", "SKU-level available stock, incoming stock, and on-hand inventory when available.")
    if any(k.startswith("catalogue_") for k in result.detected_columns):
        add_key_value_paragraph(doc, "Optional Catalogue", "Catalogue fields were linked by SKU and added to the final analysis, including category, lifecycle status, pricing, case pack, margin, and MOQ risk when available.")
    if result.purchase_summary is not None:
        add_key_value_paragraph(doc, "Optional Purchase / PO History", "Purchase amount was summarized by year for 2024, 2025, and 2026 year-to-date.")
    if result.sales_year_summary is not None:
        add_key_value_paragraph(doc, "Year-Labeled Sales Reports", "Sales files were uploaded by year, so sales trend is compared against PO purchase amount by the same uploaded year.")
    if result.location_year_business_view is not None:
        add_key_value_paragraph(doc, "Location Logic", "When Location is available in Sales, Stock, or PO files, the report keeps the main SKU analysis unchanged and adds location-level views automatically.")

    doc.add_heading("Excel Column Explanation", level=1)
    explanation = pd.DataFrame(
        [
            ["Qty", "Total SKU quantity sold in the uploaded sales period"],
            ["Contribution %", "SKU Qty / Total Qty"],
            ["Cumulative %", "Running contribution after sorting by Qty"],
            ["SABC Type", "S <= 5%, A <= 80%, B <= 95%, C > 95%"],
            ["Future Inventory", "Available + Incoming"],
            ["Adjusted Future Inventory", "MAX(0, Future Inventory)"],
            ["Avg Monthly Sales", "Qty / 12"],
            ["Coverage", "Adjusted Future Inventory / Avg Monthly Sales"],
            ["Inventory Status", "Stockout, Urgent, Healthy, Monitor, Overstock, or No Sales / Review"],
            ["Action", "Operational recommendation linked to inventory status"],
            ["2024 Purchase Amount", "PO line purchase amount matched to this SKU from the 2024 uploaded PO file"],
            ["2025 Purchase Amount", "PO line purchase amount matched to this SKU from the 2025 uploaded PO file"],
            ["2026 Purchase Amount", "PO line purchase amount matched to this SKU from the 2026 YTD uploaded PO file"],
            ["Total Purchase Amount", "Combined matched purchase amount across uploaded PO files"],
            ["Lifecycle Type", "Normalized catalogue recommendation status, such as New / Coming Soon, Discontinued, Display Risk, or Renewal"],
            ["Lifecycle Action", "Business action override driven by catalogue recommendation when applicable"],
            ["Lifecycle Note", "Explanation of why catalogue recommendation changes the SKU strategy"],
        ],
        columns=["Column", "Meaning"],
    )
    add_table(doc, explanation, max_rows=20)

    doc.add_heading("SABC Sales Analysis", level=1)
    doc.add_paragraph("The SABC view separates high-impact SKUs from long-tail products, so replenishment can protect sales while keeping slow movers controlled.")
    add_table(doc, result.sabc_summary, max_rows=10)
    sabc_png = make_sabc_chart_png(result.sabc_summary)
    doc.add_picture(sabc_png, width=Inches(6.5))

    if result.sales_purchase_year_summary is not None and not result.sales_purchase_year_summary.empty:
        doc.add_heading("Sales vs Purchase Trend by Year", level=1)
        doc.add_paragraph(
            "This view compares uploaded sales files and PO files by the same year label. It helps identify whether purchase investment is moving in line with sales demand."
        )
        add_table(doc, result.sales_purchase_year_summary, max_rows=10)
    elif result.sales_year_summary is not None and not result.sales_year_summary.empty:
        doc.add_heading("Sales Trend by Year", level=1)
        add_table(doc, result.sales_year_summary, max_rows=10)

    doc.add_heading("Inventory Coverage Analysis", level=1)
    doc.add_paragraph("Coverage translates stock into months of demand. Low coverage creates service risk; excessive coverage creates cash and warehouse pressure.")
    add_table(doc, result.inventory_status_summary, max_rows=10)
    status_png = make_bar_chart_png(result.inventory_status_summary, "Inventory Status", "SKU_Count", "Inventory Status by SKU Count")
    doc.add_picture(status_png, width=Inches(6.5))

    if result.purchase_summary is not None:
        doc.add_heading("Purchase Amount by Year", level=1)
        doc.add_paragraph(
            "This section summarizes total purchase amount by calendar year. 2026 is calculated as year-to-date based on the uploaded purchase dates."
        )
        add_table(doc, result.purchase_summary, max_rows=10)
        if result.purchase_sku_summary is not None and not result.purchase_sku_summary.empty:
            doc.add_heading("Purchase Amount by SKU", level=2)
            doc.add_paragraph("This table ranks purchase lines by SKU/Product after applying the purchase keyword filter when provided.")
            add_table(doc, result.purchase_sku_summary, max_rows=15)
            integrated_purchase_cols = [
                "Product SKU",
                "Product Name",
                "Qty",
                "SABC Type",
                "Inventory Status",
                "2024 Purchase Amount",
                "2025 Purchase Amount",
                "2026 Purchase Amount",
                "Total Purchase Amount",
            ]
            integrated_purchase_cols = [col for col in integrated_purchase_cols if col in result.final.columns]
            integrated_purchase = result.final[result.final.get("Total Purchase Amount", 0) > 0].sort_values("Total Purchase Amount", ascending=False)
            if not integrated_purchase.empty:
                doc.add_heading("Integrated Sales / Inventory / Purchase View", level=2)
                doc.add_paragraph("This view connects matched PO purchase amount back to the main SKU analysis, so purchase investment can be compared against sales and inventory status.")
                add_table(doc, integrated_purchase[integrated_purchase_cols], max_rows=15)
        if result.purchase_location_summary is not None and not result.purchase_location_summary.empty:
            doc.add_heading("Purchase Amount by Location", level=2)
            doc.add_paragraph("This view separates purchase amount by receiving location, so PO investment can be reviewed without manually filtering location before upload.")
            add_table(doc, result.purchase_location_summary, max_rows=10)
        if result.purchase_sku_location_summary is not None and not result.purchase_sku_location_summary.empty:
            doc.add_heading("Top SKU Purchase by Location", level=2)
            doc.add_paragraph("This table shows the largest SKU/location purchase combinations after applying the brand keyword filter.")
            add_table(doc, result.purchase_sku_location_summary, max_rows=15)

    if result.location_year_business_view is not None and not result.location_year_business_view.empty:
        doc.add_heading("Year + Location Business View", level=1)
        doc.add_paragraph(
            "This view is designed for questions like 2024 Toronto, 2025 Vancouver, or 2026 year-to-date by location. "
            "Sales and PO amounts are matched by year and location; stock is current inventory by location when the stock report includes Location."
        )
        add_table(doc, result.location_year_business_view, max_rows=20)
    if result.sales_location_summary is not None and not result.sales_location_summary.empty:
        doc.add_heading("Sales by Location", level=2)
        doc.add_paragraph("This table uses the uploaded sales period and separates sales quantity by location when the Sales Report includes Location.")
        add_table(doc, result.sales_location_summary, max_rows=15)
    if result.stock_location_summary is not None and not result.stock_location_summary.empty:
        doc.add_heading("Current Stock by Location", level=2)
        doc.add_paragraph("This table separates current available, incoming, on-hand, and future inventory by location when the Stock Levels Report includes Location.")
        add_table(doc, result.stock_location_summary, max_rows=15)

    if has_catalogue:
        add_catalogue_report_sections(doc, insights)

    doc.add_heading("Replenishment Priority", level=1)
    priority_cols = ["Product SKU", "Product Name", "Qty", "SABC Type", "Adjusted Future Inventory", "Coverage", "Inventory Status", "Action"]
    add_table(doc, insights["urgent_table"][priority_cols], max_rows=10)

    doc.add_heading("Overstock Risk", level=1)
    add_table(doc, insights["overstock_table"][priority_cols], max_rows=10)

    doc.add_heading("Future 6-Month Business Plan", level=1)
    add_bullet(doc, "Month 1-2: protect S/A SKUs, solve Stockout and Urgent items first, and validate incoming PO timing.")
    add_bullet(doc, "Month 2-3: reduce or pause PO for Overstock SKUs, especially C-type long-tail products.")
    add_bullet(doc, "Month 3-4: build channel-specific bundles around best-selling product lines and display sets.")
    add_bullet(doc, "Month 4-6: review SKU lifecycle, discontinue low-sales SKUs where sell-through does not improve, and reallocate budget to hero SKUs.")
    if has_catalogue:
        add_bullet(doc, "Catalogue-enhanced plan: combine sales velocity with price, margin, category, and case-pack data before finalizing replenishment, discontinuation, and new product launch decisions.")
        add_bullet(doc, "Lifecycle plan: separate new launch SKUs, discontinued SKUs, display risk SKUs, and renewal SKUs before issuing POs so replenishment does not conflict with the catalogue recommendation.")

    doc.add_heading("Label / Packaging Optimization", level=1)
    add_bullet(doc, "Prioritize bilingual label clarity for S/A products first, because these SKUs carry the highest sales impact.")
    add_bullet(doc, "Standardize shade names, claims, and barcode visibility to reduce store and warehouse handling friction.")
    add_bullet(doc, "Example: for lip tint or blush SKUs, keep the shade number, shade name, barcode, and English/French product type in the same visual area so store staff can identify variants quickly.")
    add_bullet(doc, "Example: add a small shelf-ready shade sticker or color band on outer cartons for fast warehouse picking and retail replenishment.")
    add_bullet(doc, "Example: create one consistent Canadian compliance label template for ingredients, distributor information, net content, and caution text, then apply it first to S/A SKUs.")

    doc.add_heading("Channel Growth Plan", level=1)
    add_bullet(doc, "Use S/A SKUs as traffic drivers for marketplace, boutique retail, and social commerce campaigns.")
    add_bullet(doc, "Use B SKUs for curated sets and seasonal campaigns; keep C SKUs under tight inventory review.")

    doc.add_heading("New Product Plan", level=1)
    add_bullet(doc, "Expand adjacent shades or formats only where the current line has proven sell-through.")
    add_bullet(doc, "Require catalogue pricing, MOQ, inner/outer case data, and expected launch channel before committing to new SKUs.")
    if has_catalogue:
        add_bullet(doc, "Use category and margin data to prioritize launches in product families that already show high sales contribution and healthy margin.")
        add_bullet(doc, "Avoid launching new SKUs with large outer-case quantities unless expected monthly demand can absorb the MOQ within a reasonable coverage window.")
        add_bullet(doc, "For NEW / COMING SOON items, judge them by launch readiness, display support, price/margin, and channel fit instead of historical sales alone.")
        add_bullet(doc, "For Renewal items, define a transition window: stop heavy PO on the old version, prepare the renewed SKU, and prevent duplicate overstock.")

    doc.add_heading("Manufacturer Support Needed", level=1)
    add_bullet(doc, "Confirm replenishment lead time, MOQ flexibility, tester/display support, and launch assets for priority SKUs.")
    add_bullet(doc, "Request packaging files and ingredient/claim documentation for faster Canadian channel onboarding.")
    if has_catalogue:
        add_bullet(doc, "Request support on MOQ breaks, case-pack flexibility, and wholesale pricing for SKUs with high demand but low margin or high case-pack risk.")
        add_bullet(doc, "For display low/no stock items, request tester/display replenishment separately from sellable unit replenishment, especially for S/A SKUs.")
        add_bullet(doc, "For discontinued SKUs that still have sales demand, ask the manufacturer for replacement SKU mapping or final-buy availability.")

    doc.add_heading("Key Business Conclusion", level=1)
    doc.add_paragraph(
        f"{brand_name} should manage the next 6 months with a focused SKU strategy: protect S/A winners, fix urgent replenishment gaps, "
        "and actively reduce long-tail overstock exposure. The business plan should stay data-led, with monthly updates as sales and stock reports refresh."
    )

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_outputs(
    sales_file: str | Path | BinaryIO | BytesIO | list[str | Path | BinaryIO | BytesIO],
    stock_file: str | Path | BinaryIO | BytesIO,
    catalogue_file: str | Path | BinaryIO | BytesIO | None = None,
    purchase_file: str | Path | BinaryIO | BytesIO | list[str | Path | BinaryIO | BytesIO] | None = None,
    purchase_filter_keyword: str = "",
    purchase_years: list[int] | None = None,
    brand_name: str = "Brand",
    sales_years: list[int] | None = None,
    location_filter: str = "",
) -> tuple[AnalysisResult, bytes, bytes]:
    result = build_analysis(sales_file, stock_file, catalogue_file, purchase_file, purchase_filter_keyword, purchase_years, sales_years, location_filter)
    return result, generate_excel(result), generate_word_report(result, brand_name=brand_name)
